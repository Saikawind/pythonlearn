import inspect
import os
import time
import win32com.client
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate, InlineImage
from docx import Document
import draw_torsional

# 读取模板 docx 文件
document_1 = Document("docTemplate/XXXXXXX静扭试验报告.docx")
# 中文字体
font_name = u"宋体"


# 在指定段落后插入表格
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# ---------------------------------------------------------------
# 需要在 qt 里填充的变量
# 封面
top = "test"
# 实验编号和名称
ex_id = "test"
ex_name = "test"
# 实验表格信息
# 产品型号
model_num = "test"
# 送达时间
send_time = "test"
# 中心样品编号
center_sample_id = "test"
# 送达者
sender = "test"
# 样品数量
sample_num = "test"
# 试验类型
ex_type = "test"
# 委托单位
requester = "test"
# 试验时间
ex_time = "test"
# 试验人员
ex_member = "test"
# 试验依据
ex_basis = "test"
# 试验项目
ex_project = "test"
# 试验条件
ex_conditions = "test"
# 实验图片数量
img_num = 15
# 设置实验数量，n个实验
set_ex_num = 2
# 实验种类
set_ex_type_list = ['330', '2000']
# 试验现象
set_ex_phenomenon = ["无相对滑移", "约800N·m开始滑移", "无相对滑移", "约661N·m开始滑移"]
# 试验时间
set_ex_time = ['2021/9/28', '2021/9/29']
# 试验备注
set_ex_note = ['note', 'test']
# 试验曲线图示
set_ex_describe = ['样品扭矩随转角变化曲线', '样品扭矩转角随时间变化曲线']
# 划分数据范围
annotate_data_start_list = [[0.252, 130.855], [0.558, 308.7261], [0.288, 146.8931], [0.783, 378.7886]]
annotate_data_end_list = [[0.405, 223.0181], [1.143, 640.7573], [0.423, 226.8491], [1.197, 605.229]]
# 是否绘制趋势线
predict_judge_list = [True, True, True, True]
# 预测数值范围比例，此值越大趋势线越长
predict_factor_list = [0.15, 0.15, 0.15, 0.15]
# 本地测试变量
save_time = time.strftime("%Y年%m月%d日%H-%M-%S")
save_path = "final_doc/{}静扭试验报告.docx".format(save_time)

# ---------------------------------------------------------------
# 生成试验照片动态表格
# 判断表格行数
if img_num % 3 == 0:
    row_num = int(img_num / 3)
else:
    row_num = int(img_num / 3) + 1
    # row_mod = img_num % 3

# 确定插入表格位置
index = 0
paragraph_index_con = 0
for paragraph in document_1.paragraphs:
    if paragraph.text == "试验照片":
        paragraph_index_con = index
        break
    index += 1

col_num = 4
# 实验前照片
for ex_num in range(1, set_ex_num + 1):
    table = document_1.add_table(rows=row_num, cols=col_num, style='Table Grid')
    table.autofit = False  # 很重要！
    table.cell(0, 0).merge(table.cell(row_num - 1, 0))
    table_column = table.columns[0].cells
    # 设置列宽度
    table_column[0].width = Cm(1.24)
    table_column[0].text = "{{ex_num_" + "00" + str(ex_num) + "}} 实验前照片"
    table_column[0].paragraphs[0].runs[0].font.name = font_name
    table_column[0].paragraphs[0].runs[0].font.size = Pt(15)
    img_index = 1
    # 第一列垂直居中
    table_column[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(4)
        for col in range(1, 4):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}".format(ex_num, img_index) + "}}"
            img_index += 1
    move_table_after(table, document_1.paragraphs[paragraph_index_con])
    # 添加分页符
    paragraph_index_con += 1
    new_page = document_1.add_page_break()
    document_1.paragraphs[paragraph_index_con]._p.addnext(new_page._p)
    # 添加换行
    # paragraph = document_1.add_paragraph("")
    # document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph._p)

# 实验后照片
for ex_num in range(1, set_ex_num + 1):
    for ex_type in set_ex_type_list:
        document_1.add_paragraph("")
        table = document_1.add_table(rows=row_num, cols=4, style='Table Grid')
        table.autofit = False  # 很重要！
        table.cell(0, 0).merge(table.cell(row_num - 1, 0))
        table_column = table.columns[0].cells
        # 设置列宽度
        table_column[0].width = Cm(1.24)
        table_column[0].text = "{{ex_num_" + "00" + str(ex_num) + "}}实验后照片"
        # 设置字体
        table_column[0].paragraphs[0].runs[0].font.name = font_name
        table_column[0].paragraphs[0].runs[0].font.size = Pt(15)
        # 类型字体分别设置
        table_column[0].paragraphs[0].add_run("\n" + ex_type + "N·m")
        table_column[0].paragraphs[0].runs[1].font.size = Pt(9)
        img_index = 1
        # 第一列垂直居中
        table_column[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in range(row_num):
            # 设置列高度
            table.rows[row].height = Cm(4)
            for col in range(1, 4):
                table_row = table.rows[row].cells
                # 设置图片表格宽高度
                table_row[col].width = Cm(5)
                # 垂直居中
                table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # img_1_330_1
                table_row[col].text = "{{" + "img_{}_{}_{}".format(ex_num, ex_type, img_index) + "}}"
                img_index += 1
        move_table_after(table, document_1.paragraphs[paragraph_index_con])
        # 添加分页符
        paragraph_index_con += 1
        new_page = document_1.add_page_break()
        document_1.paragraphs[paragraph_index_con]._p.addnext(new_page._p)
        # paragraph = document_1.add_paragraph("")
        # document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph._p)

# 删除多余的分页符
document_1.paragraphs[paragraph_index_con].clear()

# ---------------------------------------------------------------
# 生成试验曲线表格
# 确定表格插入位置
index = 0
paragraph_index_cur = 0
for paragraph in document_1.paragraphs:
    if paragraph.text == "试验曲线":
        paragraph_index_cur = index
        break
    index += 1

# 添加分页符
document_1.paragraphs[paragraph_index_cur - 1]._p.addnext(document_1.add_page_break()._p)
paragraph_index_cur += 1

col_num = 2
row_num = 4
describe_index = 0
for ex_num in range(set_ex_num):
    for ex_type in set_ex_type_list:
        img_index = 1
        table = document_1.add_table(rows=row_num, cols=col_num, style='Table Grid')
        # 合并单元格
        table.cell(0, 0).merge(table.cell(3, 0))
        # 取消表格自适应
        table.autofit = False
        # 第一行
        table_rows = table.rows[0].cells
        # 第一列
        table_rows[0].width = Cm(1.06)
        # 第二列
        table_rows[1].width = Cm(15.42)
        # 设置第一列文本和格式
        table.columns[0].cells[0].text = "试\n验\n曲\n线"
        table.columns[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 取第二列cell
        table_columns = table.columns[1].cells
        # 遍历第二列所有行
        describe_index = 0
        row_index = 1
        for cell in table_columns:
            # 判断行索引来填充文字还是图片
            if row_index % 2 == 0:
                cell.text = "{{curve_img_" + str(ex_num + 1) + "_" + ex_type + "_" + str(img_index) + "}}"
                img_index += 1
            elif row_index % 2 == 1:
                cell.text = "{{ex_num_00" + str(ex_num + 1) + "}}" + set_ex_describe[
                    describe_index] + "（{{ex_type_" + ex_type + "}}N·m）"
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
                describe_index += 1
            row_index += 1
            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        move_table_after(table, document_1.paragraphs[paragraph_index_cur])
        # paragraph_index_cur += 1
        # paragraph = document_1.add_paragraph("")
        # document_1.paragraphs[paragraph_index_cur]._p.addnext(paragraph._p)
        # 添加分页符
        paragraph_index_cur += 1
        new_page = document_1.add_page_break()
        document_1.paragraphs[paragraph_index_cur]._p.addnext(new_page._p)
# 删除多余的分页符
document_1.paragraphs[paragraph_index_cur].clear()

# ---------------------------------------------------------------
# 生成实验结果统计表格
table_head = 2
row_num = table_head + set_ex_num * len(set_ex_type_list)
col_num = 5
# 定位表格插入模板位置
table_last_index = len(document_1.tables) - 1
row = document_1.tables[table_last_index].rows[0].cells
row_cell = row[1]
# 生成表格
table = row_cell.add_table(rows=row_num, cols=col_num)
table.style = 'Table Grid'

# 设置表头
title_text = "定子静扭信息统计表"
table.cell(0, 0).merge(table.cell(0, 4))
table.rows[0].cells[0].text = title_text
table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(18)

# 设置表头
head_text = ["样品编号", "静扭（N·m）", "时间", "现象", "备注"]
table_row = table.rows[1].cells
index = 0
# 填充表头内容
for cell in table_row:
    cell.text = head_text[index]
    index += 1
# 设置表头高度
for cell in table_row:
    cell.height = Cm(1.19)
    for paragraph in cell.paragraphs:
        paragraph.runs[0].font.size = Pt(11)

# 表格合并行-起始
merge_before = 2
for ex_num in range(set_ex_num):
    # 表格合并行-末尾
    merge_after = merge_before + len(set_ex_type_list) - 1
    # 样品编号列
    table.cell(merge_before, 0).merge(table.cell(merge_after, 0))
    # 时间列
    table.cell(merge_before, 2).merge(table.cell(merge_after, 2))
    # 备注列
    table.cell(merge_before, 4).merge(table.cell(merge_after, 4))
    # 实验现象可能要加个判断是否合并，待定，en版本添加了
    merge_before += len(set_ex_type_list)

now_row = 1
# 该样品编号下实验类型数量，即占几行
for ex_num in range(1, set_ex_num + 1):
    # 遍历实验类型数量的行数赋值
    for ex_type in set_ex_type_list:
        now_row += 1
        row = table.rows[now_row].cells
        row[0].text = "{{ex_static_id_" + str(ex_num) + "}}"
        row[1].text = "{{ex_type_" + ex_type + "}}N·m"
        row[2].text = "{{ex_time_" + str(ex_num) + "}}"
        row[3].text = "{{ex_phenomenon_" + str(ex_num) + "_" + ex_type + "}}"
        row[4].text = "{{ex_note_" + str(ex_num) + "}}"

# 设置表格整体垂直居中
# 第几行
for i in range(row_num):
    # 第几列
    for j in range(col_num):
        row = table.rows[i].cells
        row[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc_temp_path = "generate_doc/{}test.docx".format(save_time)
# 保存生成的模板文件
document_1.save(doc_temp_path)

# ---------------------------------------------------------------

# docxtpl 以上面生成的 docx 为模板进行变量填充
doc = DocxTemplate(doc_temp_path)

# 样品安装图片，没有写扩展
install_img = InlineImage(doc, image_descriptor="data/img/install_1.png", width=Cm(16.46), height=Cm(10.95))
context = {
    'top': top,
    'ex_id': ex_id,
    'ex_name': ex_name,
    'model_num': model_num,
    'send_time': send_time,
    'center_sample_id': center_sample_id,
    'sender': sender,
    'sample_num': sample_num,
    'ex_type': ex_type,
    'requester': requester,
    'ex_time': ex_time,
    'ex_member': ex_member,
    'ex_basis': ex_basis,
    'ex_project': ex_project,
    'install_img': install_img,
    'ex_conditions': ex_conditions
}

# ---------------------------------------------------------------
# 填充图片-实验前
for ex_num in range(set_ex_num):
    for img in range(img_num):
        img_path = 'data/img/试验前照片/{}#/{}.JPG'.format(ex_num + 1, img + 1)
        context["img_{}_{}".format(ex_num + 1, img + 1)] = InlineImage(doc, image_descriptor=img_path, width=Cm(4.62),
                                                                       height=Cm(3.08))

# 实验后
for ex_type in set_ex_type_list:
    for ex_num in range(set_ex_num):
        for img in range(img_num):
            img_path = 'data/img/试验后照片/{}#/{}N 试验后照片/{}.JPG'.format(ex_num + 1, ex_type, img + 1)
            context["img_{}_{}_{}".format(ex_num + 1, ex_type, img + 1)] = InlineImage(doc, image_descriptor=img_path,
                                                                                       width=Cm(4.62),
                                                                                       height=Cm(3.08))
# ---------------------------------------------------------------
# 试验曲线表格
index = 1
# 调用绘制曲线函数
draw_torsional.draw_curve(set_ex_num, set_ex_type_list, predict_judge_list, predict_factor_list,
                          annotate_data_start_list,
                          annotate_data_end_list)
save_img_time = draw_torsional.save_img_time
for ex_num in range(1, set_ex_num + 1):
    for ex_type in set_ex_type_list:
        for img_index in range(1, 3):
            img_path = 'data/curve-{}/{}.png'.format(save_img_time, index)
            context["curve_img_{}_{}_{}".format(ex_num, ex_type, img_index)] = InlineImage(doc,
                                                                                           image_descriptor=img_path,
                                                                                           width=Cm(15.04),
                                                                                           height=Cm(8.98))
            index += 1
        context["ex_num_00{}".format(ex_num)] = "00" + str(ex_num)
        context["ex_type_{}".format(ex_type)] = ex_type

# ---------------------------------------------------------------
# 实验运行记录表格
index = 0
set_ex_num = 2
for ex_num in range(set_ex_num):
    ex_num_dict = ex_num + 1
    context["ex_static_id_" + str(ex_num_dict)] = str(ex_num_dict) + "#"
    context["ex_time_" + str(ex_num_dict)] = set_ex_time[ex_num]
    context["ex_note_" + str(ex_num_dict)] = set_ex_note[ex_num]
    for ex_type in set_ex_type_list:
        context["ex_type_" + ex_type] = ex_type
        context["ex_phenomenon_" + str(ex_num_dict) + "_" + ex_type] = set_ex_phenomenon[index]
        index += 1


# ---------------------------------------------------------------
# 目录更新
def toc_main(docx_file):
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    file_name = docx_file
    file_path = os.path.join(script_dir, file_name)
    update_toc(file_path)


def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    try:
        doc = word.Documents.Open(docx_file)
        for contentCount in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(contentCount).Update()
        for FigureCount in range(1, doc.TablesOfFigures.Count + 1):
            doc.TablesOfFigures(FigureCount).Update()
        doc.Close(SaveChanges=True)
    finally:
        word.Quit()


toc_main(doc_temp_path)
# ---------------------------------------------------------------
doc.render(context)
doc.save(save_path)
