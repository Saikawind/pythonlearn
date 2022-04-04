import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document("test.docx")

# 中文字体
font_name = u"宋体"
# table_column[0].paragraphs[0].runs[0].font.name = font_name
# table_column[0].paragraphs[0].runs[0].font.size = Pt(14)

set_ex_num = 2
set_ex_type_list = ['330', '2000']
# set_ex_type_list = ['545_10', '-545_10', '545_100', '-545_100', '1000_10', '-1000_10', '1000_100', '-1000_100',
#                     '1500_100', '-1500_100', '1900_100', '-1900_100']
col_num = 2
row_num = 4

row_index = 1
for ex_num in range(set_ex_num):
    for ex_type in set_ex_type_list:
        #
        img_index = 1
        #
        table = document.add_table(rows=row_num, cols=col_num, style='Table Grid')
        # 合并单元格
        table.cell(0, 0).merge(table.cell(3, 0))
        # 取消表格自适应
        table.autofit = False
        # 第一行
        table_rows = table.rows[0].cells
        # 第一列
        table_rows[0].width = Cm(1.06)
        # 第二列
        table_rows[1].width = Cm(15.42)
        # 设置第一列文本和格式
        table.columns[0].cells[0].text = "试\n验\n曲\n线"
        table.columns[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 取第二列cell
        table_columns = table.columns[1].cells
        # 遍历第二列所有行
        for cell in table_columns:
            # 判断行索引来填充文字还是图片
            if row_index % 2 == 0:
                cell.text = "{{curve_img_" + str(ex_num + 1) + "_" + ex_type + "_" + str(img_index) + "}}"
                img_index += 1
            else:
                cell.text = "{{ex_num_00" + str(ex_num + 1) + "}} 样品扭矩随转角变化曲线（{{ex_type_" + ex_type + "}}N·m）"
            row_index += 1
            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        document.add_paragraph("")

now_time = time.strftime("%Y年%m月%d日%H-%M-%S")
document.save("generate_doc/{}test.docx".format(now_time))

# --------------------------------------------------------

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}
index = 1
for ex_num in range(1, set_ex_num + 1):
    for ex_type in set_ex_type_list:
        for img_index in range(1, 3):
            img_path = 'data/curve_img/{}.png'.format(index)
            context["curve_img_{}_{}_{}".format(ex_num, ex_type, img_index)] = InlineImage(doc,
                                                                                           image_descriptor=img_path,
                                                                                           width=Cm(15.04),
                                                                                           height=Cm(8.98))
            index += 1
        context["ex_num_00{}".format(ex_num)] = "00" + str(ex_num)
        context["ex_type_{}".format(ex_type)] = ex_type

doc.render(context)
doc.save("doc/{}test.docx".format(now_time))
