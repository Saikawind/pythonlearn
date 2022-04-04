import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document("test.docx")

img_num = 15
# 取余
if img_num % 3 == 0:
    row_num = int(img_num / 3)
else:
    row_num = int(img_num / 3) + 1
    # row_mod = img_num % 3

# 中文字体
font_name = u"宋体"

# 设置实验数量，n个实验
set_ex_num = 2
# 实验前照片
for ex_num in range(1, set_ex_num + 1):
    table = document.add_table(rows=row_num, cols=4, style='Table Grid')
    table.autofit = False  # 很重要！
    table.cell(0, 0).merge(table.cell(row_num - 1, 0))
    table_column = table.columns[0].cells
    # 设置列宽度
    table_column[0].width = Cm(1.24)
    table_column[0].text = "{{ex_num_" + "00" + str(ex_num) + "}} 实验前照片"
    table_column[0].paragraphs[0].runs[0].font.name = font_name
    table_column[0].paragraphs[0].runs[0].font.size = Pt(14)
    # run = table_column[0].paragraphs[0].add_run("{{ex_num_" + "00" + str(ex_num) + "}} 实验前照片")
    # run.font.name = font_name
    # run._element.rPr.rFonts.set(qn('ct:eastAsia'), font_name)
    img_index = 1
    # 第一列垂直居中
    table_column[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(4)
        for col in range(1, 4):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}".format(ex_num, img_index) + "}}"
            img_index += 1
    # 添加换行
    document.add_paragraph("")

# 实验后照片
set_ex_type_list = ['330', '2000']
for ex_num in range(1, set_ex_num + 1):
    for ex_type in set_ex_type_list:
        table = document.add_table(rows=row_num, cols=4, style='Table Grid')
        table.autofit = False  # 很重要！
        table.cell(0, 0).merge(table.cell(row_num - 1, 0))
        table_column = table.columns[0].cells
        # 设置列宽度
        table_column[0].width = Cm(1.24)
        table_column[0].text = "{{ex_num_" + "00" + str(ex_num) + "}}实验后照片"
        # 设置字体
        table_column[0].paragraphs[0].runs[0].font.name = font_name
        table_column[0].paragraphs[0].runs[0].font.size = Pt(14)
        # 类型字体分别设置
        table_column[0].paragraphs[0].add_run(ex_type + "N·m")
        table_column[0].paragraphs[0].runs[1].font.size = Pt(12)

        img_index = 1
        # 第一列垂直居中
        table_column[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in range(row_num):
            # 设置列高度
            table.rows[row].height = Cm(4)
            for col in range(1, 4):
                table_row = table.rows[row].cells
                # 设置图片表格宽高度
                table_row[col].width = Cm(5)
                # 垂直居中
                table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # img_1_330_1
                table_row[col].text = "{{" + "img_{}_{}_{}".format(ex_num, ex_type, img_index) + "}}"
                img_index += 1
        # 添加换行
        document.add_paragraph("")

now_time = time.strftime("%Y年%m月%d日%H-%M-%S")
document.save("generate_doc/{}test.docx".format(now_time))

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}

for i in range(1, set_ex_num + 1):
    ex_num = '00' + str(i)
    context['ex_num_' + '00' + str(i)] = ex_num

# 填充图片-实验前
for ex_num in range(1, set_ex_num + 1):
    for img in range(1, img_num + 1):
        img_path = 'data/img/试验前照片/{}#/{}.JPG'.format(ex_num, img)
        context["img_{}_{}".format(ex_num, img)] = InlineImage(doc, image_descriptor=img_path, width=Cm(4.62),
                                                               height=Cm(3.08))

# 实验后
for ex_type in set_ex_type_list:
    for ex_num in range(1, set_ex_num + 1):
        for img in range(1, img_num + 1):
            img_path = 'data/img/试验后照片/{}#/{}N 试验后照片/{}.JPG'.format(ex_num, ex_type, img)
            context["img_{}_{}_{}".format(ex_num, ex_type, img)] = InlineImage(doc, image_descriptor=img_path,
                                                                               width=Cm(4.62),
                                                                               height=Cm(3.08))

doc.render(context)

doc.save("doc/{}test.docx".format(now_time))
