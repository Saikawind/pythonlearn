import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document("test.docx")

# 中文字体
font_name = u"宋体"
# table_column[0].paragraphs[0].runs[0].font.name = font_name
# table_column[0].paragraphs[0].runs[0].font.size = Pt(14)


# set_ex_type_list = ['545N·m_10', '-545_10N·m_10', '545_100N·m_10', '-545_100N·m_10', '1000_10N·m_10',
# ' -1000_10N·m_10', '1000_100N·m_10', '-1000_100N·m_10', '1500_100N·m_10', '-1500_100N·m_10', '1900_100N·m_10',
# '-1900_100N·m_10']

# 负号(-)有匹配问题，所以换掉负号
set_ex_type_list_transfer = ['545_10', 'minus_545_10', '545_100', 'minus_545_100', '1000_10', 'minus_1000_10',
                             '1000_100',
                             'minus_1000_100',
                             '1500_100', 'minus_1500_100', '1900_100', 'minus_1900_100']
# 设置实验现象表格是否合并
phenomenon_merge = True

set_ex_num = 1

# 表头所占行数
table_head = 2
row_num = table_head + set_ex_num * len(set_ex_type_list_transfer)
col_num = 5
table = document.add_table(rows=row_num, cols=col_num, style='Table Grid')

# 表头设置
title_text = "定子静扭信息统计表"
table.cell(0, 0).merge(table.cell(0, 4))
table.rows[0].cells[0].text = title_text
head_text = ["样品编号\nnumber", "静扭（N·m）", "时间\ntime", "现象\nphenomenon", "备注\nremark"]
table_row = table.rows[1].cells
index = 0
for cell in table_row:
    cell.text = head_text[index]
    index += 1

# 设置第二列宽度
row = table.rows[2].cells
first_col = row[1].width = Cm(3.5)

merge_before = 2
for ex_num in range(set_ex_num):
    merge_after = merge_before + len(set_ex_type_list_transfer) - 1
    # 样品编号
    table.cell(merge_before, 0).merge(table.cell(merge_after, 0))
    # 时间
    table.cell(merge_before, 2).merge(table.cell(merge_after, 2))
    # 备注
    table.cell(merge_before, 4).merge(table.cell(merge_after, 4))
    # 现象
    if phenomenon_merge:
        table.cell(merge_before, 3).merge(table.cell(merge_after, 3))
    # 实验现象可能要加个判断是否合并，待定

    merge_before += len(set_ex_type_list_transfer)

now_row = 1
# 该样品编号下实验类型数量，即占几行
for ex_num in range(1, set_ex_num + 1):
    # 遍历实验类型数量的行数赋值
    for ex_type in set_ex_type_list_transfer:
        now_row += 1
        row = table.rows[now_row].cells
        row[0].text = "{{ex_static_id_" + str(ex_num) + "}}"
        row[1].text = "{{ex_type_" + ex_type + "}}N·m/s"
        row[2].text = "{{ex_time_" + str(ex_num) + "}}"
        row[3].text = "{{ex_phenomenon_" + str(ex_num) + "_" + ex_type + "}}"
        row[4].text = "{{ex_note_" + str(ex_num) + "}}"

# 第几行
for i in range(row_num):
    # 第几列
    for j in range(col_num):
        row = table.rows[i].cells
        row[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

now_time = time.strftime("%Y年%m月%d日%H-%M-%S")
document.save("generate_doc/{}test.docx".format(now_time))

# --------------------------------------------------------

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}
set_ex_time = ['2020/1/10']
set_ex_phenomenon = ["无明显相对滑移No obvious relative slip"]
set_ex_note = ['note']

# set_ex_type_list_transfer = ['545_10', 'minus_545_10', '545_100', 'minus_545_100', '1000_10', 'minus_1000_10',
#                              '1000_100',
#                              'minus_1000_100',
#                              '1500_100', 'minus_1500_100', '1900_100', 'minus_1900_100']

# 填充字符内容使用未转换的列表
set_ex_type_list = ['545_10', '-545N·m_10', '545N·m_10', '-545N·m_10', '1000N·m_10', ' -1000N·m_10',
                    '1000N·m_100', '-1000N·m_100',
                    '1500N·m_100', '-1500N·m_100', '1900N·m_100', '-1900N·m_100']

for ex_num in range(set_ex_num):
    ex_num_dict = ex_num + 1
    context["ex_static_id_" + str(ex_num_dict)] = str(ex_num_dict) + "#"
    context["ex_time_" + str(ex_num_dict)] = set_ex_time[ex_num]
    context["ex_note_" + str(ex_num_dict)] = set_ex_note[ex_num]
    for index in range(len(set_ex_type_list_transfer)):
        context["ex_type_" + set_ex_type_list_transfer[index]] = set_ex_type_list[index]
        context["ex_phenomenon_" + str(ex_num_dict) + "_" + set_ex_type_list_transfer[index]] = set_ex_phenomenon[
            ex_num]

doc.render(context)
doc.save("doc/{}test.docx".format(now_time))
