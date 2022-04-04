import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document("test.docx")

# 中文字体
font_name = u"宋体"
# table_column[0].paragraphs[0].runs[0].font.name = font_name
# table_column[0].paragraphs[0].runs[0].font.size = Pt(14)

set_ex_num = 2
set_ex_type_list = ['330', '2000']
# set_ex_type_list = ['545_10', '-545_10', '545_100', '-545_100', '1000_10', '-1000_10', '1000_100', '-1000_100',
#                     '1500_100', '-1500_100', '1900_100', '-1900_100']
# 表头所占行数
table_head = 2
row_num = table_head + set_ex_num * len(set_ex_type_list)
col_num = 5
table = document.add_table(rows=row_num, cols=col_num, style='Table Grid')

title_text = "定子静扭信息统计表"
table.cell(0, 0).merge(table.cell(0, 4))
table.rows[0].cells[0].text = title_text
head_text = ["样品编号", "静扭（N·m）", "时间", "现象", "备注"]
table_row = table.rows[1].cells
index = 0
for cell in table_row:
    cell.text = head_text[index]
    index += 1

merge_before = 2
for ex_num in range(set_ex_num):
    merge_after = merge_before + len(set_ex_type_list) - 1
    # 样品编号
    table.cell(merge_before, 0).merge(table.cell(merge_after, 0))
    # 时间
    table.cell(merge_before, 2).merge(table.cell(merge_after, 2))
    # 备注
    table.cell(merge_before, 4).merge(table.cell(merge_after, 4))
    # 实验现象可能要加个判断是否合并，待定

    merge_before += len(set_ex_type_list)

now_row = 1
# 该样品编号下实验类型数量，即占几行
for ex_num in range(1, set_ex_num + 1):
    # 遍历实验类型数量的行数赋值
    for ex_type in set_ex_type_list:
        now_row += 1
        row = table.rows[now_row].cells
        row[0].text = "{{ex_static_id_" + str(ex_num) + "}}"
        row[1].text = "{{ex_type_" + ex_type + "}}N·m"
        row[2].text = "{{ex_time_" + str(ex_num) + "}}"
        row[3].text = "{{ex_phenomenon_" + str(ex_num) + "_" + ex_type + "}}"
        row[4].text = "{{ex_note_" + str(ex_num) + "}}"

# 第几行
for i in range(row_num):
    # 第几列
    for j in range(col_num):
        row = table.rows[i].cells
        row[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

row = table.rows[5].cells[4]
table_2 = row.add_table(rows=2, cols=2)
row.tables[0].style = 'Table Grid'

now_time = time.strftime("%Y年%m月%d日%H-%M-%S")
document.save("generate_doc/{}test.docx".format(now_time))

# --------------------------------------------------------

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}
set_ex_time = ['2021/9/28', '2021/9/29']
set_ex_phenomenon = ["无相对滑移", "约800N·m开始滑移", "无相对滑移", "约661N·m开始滑移"]
set_ex_note = ['note', 'test']

index = 0
for ex_num in range(set_ex_num):
    ex_num_dict = ex_num + 1
    context["ex_static_id_" + str(ex_num_dict)] = str(ex_num_dict) + "#"
    context["ex_time_" + str(ex_num_dict)] = set_ex_time[ex_num]
    context["ex_note_" + str(ex_num_dict)] = set_ex_note[ex_num]
    for ex_type in set_ex_type_list:
        context["ex_type_" + ex_type] = ex_type
        context["ex_phenomenon_" + str(ex_num_dict) + "_" + ex_type] = set_ex_phenomenon[index]
        index += 1

doc.render(context)
doc.save("doc/{}test.docx".format(now_time))
