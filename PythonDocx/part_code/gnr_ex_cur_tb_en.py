import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

import draw_torsional_en

document = Document("test.docx")

# 中文字体
font_name = u"宋体"

set_ex_num = 1
set_ex_type_list_transfer = ['545_10', 'minus_545_10', '545_100', 'minus_545_100', '1000_10', 'minus_1000_10',
                             '1000_100',
                             'minus_1000_100',
                             '1500_100', 'minus_1500_100', '1900_100', 'minus_1900_100']
col_num = 1
row_num = 4
row_index = 1
for ex_num in range(set_ex_num):
    for ex_type in set_ex_type_list_transfer:
        #
        img_index = 1
        #
        table = document.add_table(rows=row_num, cols=col_num, style='Table Grid')
        # 取消表格自适应
        table.autofit = False
        # 取第一列cell
        table_columns = table.columns[0].cells
        # 遍历第一列所有行
        for cell in table_columns:
            # 判断行索引来填充文字还是图片
            if row_index % 2 == 0:
                cell.text = "{{curve_img_" + str(ex_num + 1) + "_" + ex_type + "_" + str(img_index) + "}}"
                img_index += 1
            else:
                cell.text = "{{ex_num_00" + str(ex_num + 1) + "}} 样品扭矩随转角变化曲线（{{ex_type_" + ex_type + "}}N·m/s）"
            row_index += 1
            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        document.add_paragraph("")

for ex_num in range(set_ex_num):
    img_index = 1
    #
    table = document.add_table(rows=set_ex_num * 2, cols=col_num, style='Table Grid')
    # 取消表格自适应
    table.autofit = False
    # 取第一列cell
    table_columns = table.columns[0].cells
    # 遍历第一列所有行
    for cell in table_columns:
        # 判断行索引来填充文字还是图片
        if row_index % 2 == 0:
            cell.text = "{{curve_img_" + str(ex_num + 1) + "_heating_" + str(img_index) + "}}"
            img_index += 1
        else:
            cell.text = "{{ex_num_00" + str(ex_num + 1) + "}} 升温曲线图\n" + "{{ex_num_00" + str(
                ex_num + 1) + "}} Heating curve"
        row_index += 1
        # 垂直居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

now_time = draw_torsional_en.save_img_time
document.save("generate_doc/{}test.docx".format(now_time))

# --------------------------------------------------------

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}
set_ex_type_list = ['545N·m_10', '-545N·m_10', '545N·m_10', '-545N·m_10', '1000N·m_10', ' -1000N·m_10',
                    '1000N·m_100', '-1000N·m_100',
                    '1500N·m_100', '-1500N·m_100', '1900N·m_100', '-1900N·m_100']
# 普通实验两张图
index = 1
draw_torsional_en.draw_curve_en()
for ex_num in range(1, set_ex_num + 1):
    for type_index in range(len(set_ex_type_list_transfer)):
        # (1, 3)索引一个表格里的两个实验曲线图
        for img_index in range(1, 3):
            img_path = 'data/curve-{}-EN/{}.png'.format(now_time, index)
            context[
                "curve_img_{}_{}_{}".format(ex_num, set_ex_type_list_transfer[type_index], img_index)] = InlineImage(
                doc,
                image_descriptor=img_path,
                width=Cm(15.04),
                height=Cm(8.98))
            index += 1
        context["ex_num_00{}".format(ex_num)] = str(ex_num) + "#"
        context["ex_type_{}".format(set_ex_type_list_transfer[type_index])] = set_ex_type_list[type_index]

# 升温曲线
img_index = 1
for ex_num in range(1, set_ex_num + 1):
    img_path = 'data/curve-{}-EN/{}.png'.format(now_time, index)
    context["curve_img_{}_heating_{}".format(ex_num, img_index)] = InlineImage(
        doc,
        image_descriptor=img_path,
        width=Cm(15.04),
        height=Cm(8.98))
    img_index += 1
    index += 1

doc.render(context)
doc.save("doc/{}test.docx".format(now_time))
