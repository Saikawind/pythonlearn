import time

import lxml
from docx import Document
from docx.opc.oxml import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

document = Document("test.docx")

img_num = 12
# 取余
if img_num % 3 == 0:
    row_num = int(img_num / 3)
else:
    row_num = int(img_num / 3) + 1
    # row_mod = img_num % 3

# 中文字体
font_name = u"宋体"

# 设置实验数量，n个实验
set_ex_num = 1
# 实验前照片
for ex_num in range(1, set_ex_num + 1):
    paragraph_1 = document.add_paragraph("{}#样品试验前照片：".format(ex_num))
    paragraph_2 = document.add_paragraph("{}# Photos before sample test:".format(ex_num))
    paragraph_1.alignment = WD_TABLE_ALIGNMENT.LEFT
    paragraph_2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table = document.add_table(rows=row_num, cols=3, style='Table Grid')
    table.autofit = False  # 很重要！
    img_index = 1
    # 第一列垂直居中
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(5.5)
        for col in range(3):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}".format(ex_num, img_index) + "}}"
            img_index += 1
    # 添加换行
    document.add_page_break()

# 实验后照片
for ex_num in range(1, set_ex_num + 1):
    paragraph_1 = document.add_paragraph("{}#样品试验后照片：".format(ex_num))
    paragraph_2 = document.add_paragraph("{}# Photos after sample test:".format(ex_num))
    paragraph_1.alignment = WD_TABLE_ALIGNMENT.LEFT
    paragraph_2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table = document.add_table(rows=row_num, cols=3, style='Table Grid')
    table.autofit = False  # 很重要！
    img_index = 1
    # 第一列垂直居中
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(5.5)
        for col in range(3):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}_after".format(ex_num, img_index) + "}}"
            img_index += 1
    # 添加换行
    document.add_page_break()

now_time = time.strftime("%Y年%m月%d日%H-%M-%S")
document.save("generate_doc/{}test.docx".format(now_time))

doc = DocxTemplate("generate_doc/{}test.docx".format(now_time))

context = {}
#
for i in range(1, set_ex_num + 1):
    ex_num = '00' + str(i)
    context['ex_num_' + '00' + str(i)] = ex_num

# 填充图片-实验前
for ex_num in range(1, set_ex_num + 1):
    for img in range(1, img_num + 1):
        img_path = 'data/img_en/试验前/{}#/{}.JPG'.format(ex_num, img)
        context["img_{}_{}".format(ex_num, img)] = InlineImage(doc, image_descriptor=img_path, width=Cm(4.61),
                                                               height=Cm(3.07))
#
# 实验后
for ex_num in range(1, set_ex_num + 1):
    for img in range(1, img_num + 1):
        img_path = 'data/img_en/试验后/{}#/{}.JPG'.format(ex_num, img)
        context["img_{}_{}_after".format(ex_num, img)] = InlineImage(doc, image_descriptor=img_path,
                                                                     width=Cm(4.61),
                                                                     height=Cm(3.07))

doc.render(context)

doc.save("doc/{}test.docx".format(now_time))
