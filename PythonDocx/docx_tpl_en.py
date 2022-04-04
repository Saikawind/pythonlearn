import time
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt
from docxtpl import DocxTemplate, InlineImage
from docx import Document
import draw_torsional
import draw_torsional_en
import win32com.client
import inspect, os

# 读取模板 docx 文件
document_1 = Document("docTemplate/远景定子高温静扭试验报告.docx")
# 中文字体
font_name = u"宋体"


# 在指定段落后插入表格
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# ---------------------------------------------------------------
# 需要在 qt 里填充的变量
# 页眉
top = "test"
# 实验编号和名称
ex_id = "test"
ex_name = "test"
# 实验表格信息
# 产品型号
model_num = "test"
# 送达时间
send_time = "test"
# 中心样品编号
center_sample_id = "test"
# 送达者
sender = "test"
# 样品数量
sample_num = "test"
# 试验类型
ex_type = "test"
# 委托单位
requester = "test"
# 试验时间
ex_time = "test"
# 试验人员
ex_member = "test"
# 试验依据
ex_basis = "test"
# 试验项目
ex_project = "test"
# 试验条件
ex_conditions = "test"
# 试验地点
ex_location = "test"
# 试验结论
ex_result = "test"
# 实验数量
set_ex_num = 1
# 给定图片数量
img_num = 12
# 未转换的试验种类列表
# 填充字符内容使用未转换的列表，不做 context 变量，用来填充模板内容
set_ex_type_list = ['545N·m-10', '-545N·m-10', '545N·m-10', '-545N·m-10', '1000N·m-10', '-1000N·m-10',
                    '1000N·m-100', '-1000N·m-100',
                    '1500N·m-100', '-1500N·m-100', '1900N·m-100', '-1900N·m-100']
# 试验曲线图示
set_ex_describe = ['样品扭矩随转角变化曲线', '样品扭矩转角随时间变化曲线']
set_ex_describe_en = ['Sample torque curve with Angle', 'Sample torque Angle curve with time']
# 试验现象
set_ex_phenomenon = ["无明显相对滑移No obvious relative slip"]
# 试验时间
set_ex_time = ['2020/1/10']
# 试验备注
set_ex_note = ['note']
# 划分数据范围
# annotate_data_start_list = [[0.252, 130.855], [0.558, 308.7261], [0.288, 146.8931], [0.783, 378.7886]]
# annotate_data_end_list = [[0.405, 223.0181], [1.143, 640.7573], [0.423, 226.8491], [1.197, 605.229]]
# 是否绘制趋势线
# predict_judge_list = [True, True, True, True]
# 预测数值范围比例，此值越大趋势线越长
# predict_factor_list = [0.15, 0.15, 0.15, 0.15]
# 设置实验现象表格是否合并
phenomenon_merge = True

# 本地测试变量
save_time = time.strftime("%Y年%m月%d日%H-%M-%S")
save_path = "final_doc/{}远景定子高温静扭试验报告.docx".format(save_time)
# 字符匹配转换的试验类型列表
set_ex_type_list_transfer = []
for ex_type in set_ex_type_list:
    ex_type = ex_type.replace("N·m-", "_")
    ex_type = ex_type.replace("-", "minus")
    set_ex_type_list_transfer.append(ex_type)
# set_ex_type_list_transfer = ['545_10', 'minus545_10', '545_100', 'minus545_100', '1000_10', 'minus1000_10',
#                              '1000_100',
#                              'minus1000_100',
#                              '1500_100', 'minus1500_100', '1900_100', 'minus1900_100']

# ---------------------------------------------------------------
# 生成实验照片动态表格

# 判断表格行数
if img_num % 3 == 0:
    row_num = int(img_num / 3)
else:
    row_num = int(img_num / 3) + 1

index = 0
paragraph_index_con = 0
for paragraph in document_1.paragraphs:
    if paragraph.text == "试验照片 Test photos:":
        paragraph_index_con = index
        break
    index += 1

# 设置实验数量，n个实验
set_ex_num = 1
# 实验前照片
for ex_num in range(1, set_ex_num + 1):
    paragraph_1 = document_1.add_paragraph("{}#样品试验前照片：".format(ex_num))
    document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph_1._p)
    paragraph_index_con += 1
    paragraph_2 = document_1.add_paragraph("{}# Photos before sample test:".format(ex_num))
    document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph_2._p)
    paragraph_index_con += 1
    paragraph_1.alignment = WD_TABLE_ALIGNMENT.LEFT
    paragraph_2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table = document_1.add_table(rows=row_num, cols=3, style='Table Grid')
    table.autofit = False  # 很重要！
    img_index = 1
    # 第一列垂直居中
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(5.5)
        for col in range(3):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}".format(ex_num, img_index) + "}}"
            img_index += 1
    # 添加换行
    move_table_after(table, document_1.paragraphs[paragraph_index_con])
    # 添加分页符
    paragraph_index_con += 1
    new_page = document_1.add_page_break()
    document_1.paragraphs[paragraph_index_con]._p.addnext(new_page._p)
    # 创建空行
    # paragraph = document_1.add_paragraph("")
    # 添加空行到指定位置
    # document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph._p)

# 实验后照片
for ex_num in range(1, set_ex_num + 1):
    paragraph_1 = document_1.add_paragraph("{}#样品试验后照片：".format(ex_num))
    document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph_1._p)
    paragraph_index_con += 1
    paragraph_2 = document_1.add_paragraph("{}# Photos after sample test:".format(ex_num))
    document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph_2._p)
    paragraph_index_con += 1
    paragraph_1.alignment = WD_TABLE_ALIGNMENT.LEFT
    paragraph_2.alignment = WD_TABLE_ALIGNMENT.LEFT
    table = document_1.add_table(rows=row_num, cols=3, style='Table Grid')
    table.autofit = False  # 很重要！
    img_index = 1
    # 第一列垂直居中
    for row in range(row_num):
        # 设置列高度
        table.rows[row].height = Cm(5.5)
        for col in range(3):
            table_row = table.rows[row].cells
            # 设置图片表格宽高度
            table_row[col].width = Cm(5)
            # 垂直居中
            table_row[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内填充文字
            table_row[col].text = "{{" + "img_{}_{}_after".format(ex_num, img_index) + "}}"
            img_index += 1
    # 添加换行
    move_table_after(table, document_1.paragraphs[paragraph_index_con])
    paragraph_index_con += 1
    # paragraph = document_1.add_paragraph("")
    # document_1.paragraphs[paragraph_index_con]._p.addnext(paragraph._p)
    # 添加分页符
    new_page = document_1.add_page_break()
    document_1.paragraphs[paragraph_index_con]._p.addnext(new_page._p)

# 删除多余的分页符
document_1.paragraphs[paragraph_index_con].clear()

# ---------------------------------------------------------------
# 生成试验曲线表格
# 确定表格插入位置
index = 0
paragraph_index_cur = 0
for paragraph in document_1.paragraphs:
    if paragraph.text == "试验曲线 Test curve: ":
        paragraph_index_cur = index
        break
    index += 1

# 添加分页符
document_1.paragraphs[paragraph_index_cur - 1]._p.addnext(document_1.add_page_break()._p)
paragraph_index_cur += 1

col_num = 1
row_num = 4
for ex_num in range(set_ex_num):
    for ex_type in set_ex_type_list_transfer:
        #
        img_index = 1
        #
        table = document_1.add_table(rows=row_num, cols=col_num, style='Table Grid')
        # 取消表格自适应
        table.autofit = False
        # 取第一列cell
        table_columns = table.columns[0].cells
        # 遍历第一列所有行
        describe_index = 0
        row_index = 1
        for cell in table_columns:
            # 判断行索引来填充文字还是图片
            if row_index % 2 == 0:
                cell.text = "{{curve_img_" + str(ex_num + 1) + "_" + ex_type + "_" + str(img_index) + "}}"
                img_index += 1
            elif row_index % 2 == 1:
                cell.text = "{{ex_num_00" + str(
                    ex_num + 1) + "}}" + set_ex_describe[
                                describe_index] + "（{{ex_type_" + ex_type + "}}N·m/s）\n" + "{{ex_num_00" + str(
                    ex_num + 1) + "}}" + set_ex_describe_en[describe_index] + "（{{ex_type_" + ex_type + "}}N·m/s）"
                describe_index += 1
                cell.paragraphs[0].runs[0].font.size = Pt(10.5)
            row_index += 1
            # 垂直居中
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # 表格内文字居中
            table_paragraphs = cell.paragraphs
            for paragraph in table_paragraphs:
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
        move_table_after(table, document_1.paragraphs[paragraph_index_cur])
        # paragraph = document_1.add_paragraph("")
        # document_1.paragraphs[paragraph_index_cur]._p.addnext(paragraph._p)
        # 添加分页符
        paragraph_index_cur += 1
        new_page = document_1.add_page_break()
        document_1.paragraphs[paragraph_index_cur]._p.addnext(new_page._p)

for ex_num in range(set_ex_num):
    table = document_1.add_table(rows=set_ex_num * 2, cols=col_num, style='Table Grid')
    # 取消表格自适应
    table.autofit = False
    # 取第一列cell
    table_columns = table.columns[0].cells
    # 遍历第一列所有行
    img_index = 1
    row_index = 1
    for cell in table_columns:
        # 判断行索引来填充文字还是图片
        if row_index % 2 == 0:
            cell.text = "{{curve_img_" + str(ex_num + 1) + "_heating_" + str(img_index) + "}}"
            img_index += 1
        else:
            cell.text = "{{ex_num_00" + str(ex_num + 1) + "}} 升温曲线图\n" + "{{ex_num_00" + str(
                ex_num + 1) + "}} Heating curve"
            cell.paragraphs[0].runs[0].font.size = Pt(10.5)
        row_index += 1
        # 垂直居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 表格内文字居中
        table_paragraphs = cell.paragraphs
        for paragraph in table_paragraphs:
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
    move_table_after(table, document_1.paragraphs[paragraph_index_cur])
    # paragraph = document_1.add_paragraph("")
    # paragraph_index_cur += 1
    # document_1.paragraphs[paragraph_index_cur]._p.addnext(paragraph._p)
    # 添加分页符
    paragraph_index_cur += 1
    new_page = document_1.add_page_break()
    document_1.paragraphs[paragraph_index_cur]._p.addnext(new_page._p)
# 删除多余的分页符
document_1.paragraphs[paragraph_index_cur].clear()

# ---------------------------------------------------------------
# 生成实验结果统计表格

# 表头所占行数
table_head = 2
row_num = table_head + set_ex_num * len(set_ex_type_list_transfer)
col_num = 5
# 定位表格插入模板位置
table_last_index = len(document_1.tables) - 1
row = document_1.tables[table_last_index].rows[0].cells
row_cell = row[1]
# 生成表格
table = row_cell.add_table(rows=row_num, cols=col_num)
table.style = 'Table Grid'

# 表头设置
title_text = "定子静扭信息统计表"
table.cell(0, 0).merge(table.cell(0, 4))
table.rows[0].cells[0].text = title_text
table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(18)

# 表头设置
head_text = ["样品编号\nnumber", "静扭（N·m）", "时间\ntime", "现象\nphenomenon", "备注\nremark"]
table_row = table.rows[1].cells
index = 0
for cell in table_row:
    cell.text = head_text[index]
    index += 1
# 设置表头高度
for cell in table_row:
    cell.height = Cm(1.19)
    for paragraph in cell.paragraphs:
        paragraph.runs[0].font.size = Pt(11)

# 设置第二列宽度
row = table.rows[2].cells
first_col = row[1].width = Cm(4)

merge_before = 2
for ex_num in range(set_ex_num):
    merge_after = merge_before + len(set_ex_type_list_transfer) - 1
    # 样品编号
    table.cell(merge_before, 0).merge(table.cell(merge_after, 0))
    # 时间
    table.cell(merge_before, 2).merge(table.cell(merge_after, 2))
    # 备注
    table.cell(merge_before, 4).merge(table.cell(merge_after, 4))
    # 现象
    if phenomenon_merge:
        table.cell(merge_before, 3).merge(table.cell(merge_after, 3))

    merge_before += len(set_ex_type_list_transfer)

now_row = 1
# 该样品编号下实验类型数量，即占几行
for ex_num in range(1, set_ex_num + 1):
    # 遍历实验类型数量的行数赋值
    for ex_type in set_ex_type_list_transfer:
        now_row += 1
        row = table.rows[now_row].cells
        row[0].text = "{{ex_static_id_" + str(ex_num) + "}}"
        row[1].text = "{{ex_type_" + ex_type + "}}N·m/s"
        row[2].text = "{{ex_time_" + str(ex_num) + "}}"
        row[3].text = "{{ex_phenomenon_" + str(ex_num) + "_" + ex_type + "}}"
        row[4].text = "{{ex_note_" + str(ex_num) + "}}"

# 第几行
for i in range(row_num):
    # 第几列
    for j in range(col_num):
        row = table.rows[i].cells
        row[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 表格内文字居中
        table_paragraphs = row[j].paragraphs
        for paragraph in table_paragraphs:
            paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER

# 保存生成的模板文件
doc_temp_path = "generate_doc/{}test.docx".format(save_time)
document_1.save(doc_temp_path)

# ---------------------------------------------------------------

# docxtpl 以上面生成的 docx 为模板进行变量填充
doc = DocxTemplate(doc_temp_path)

# sample install
# install_img = InlineImage(doc, image_descriptor="data/img/install_1.png", width=Cm(16.46), height=Cm(10.95))
context = {
    'top': top,
    'ex_id': ex_id,
    'ex_name': ex_name,
    'model_num': model_num,
    'send_time': send_time,
    'center_sample_id': center_sample_id,
    'sender': sender,
    'sample_num': sample_num,
    'ex_type': ex_type,
    'requester': requester,
    'ex_time': ex_time,
    'ex_member': ex_member,
    'ex_basis': ex_basis,
    'ex_project': ex_project,
    # 'install_img': install_img,
    'ex_conditions': ex_conditions,
    'ex_location': ex_location,
    'ex_result': ex_result
}

# ---------------------------------------------------------------
# 填充图片-实验前
for ex_num in range(1, set_ex_num + 1):
    for img in range(1, img_num + 1):
        img_path = 'data/img_en/试验前/{}#/{}.JPG'.format(ex_num, img)
        context["img_{}_{}".format(ex_num, img)] = InlineImage(doc, image_descriptor=img_path, width=Cm(4.61),
                                                               height=Cm(3.07))
#
# 实验后
for ex_num in range(1, set_ex_num + 1):
    for img in range(1, img_num + 1):
        img_path = 'data/img_en/试验后/{}#/{}.JPG'.format(ex_num, img)
        context["img_{}_{}_after".format(ex_num, img)] = InlineImage(doc, image_descriptor=img_path,
                                                                     width=Cm(4.61),
                                                                     height=Cm(3.07))

# ---------------------------------------------------------------
# 普通实验两张图
index = 1
draw_torsional_en.draw_curve_en(set_ex_num, set_ex_type_list_transfer)
save_img_time = draw_torsional_en.save_img_time
for ex_num in range(1, set_ex_num + 1):
    for type_index in range(len(set_ex_type_list_transfer)):
        # (1, 3)索引一个表格里的两个实验曲线图
        for img_index in range(1, 3):
            img_path = 'data/curve-{}-EN/{}.png'.format(save_img_time, index)
            context[
                "curve_img_{}_{}_{}".format(ex_num, set_ex_type_list_transfer[type_index], img_index)] = InlineImage(
                doc,
                image_descriptor=img_path,
                width=Cm(16.26),
                height=Cm(9.15))
            index += 1
        context["ex_num_00{}".format(ex_num)] = str(ex_num) + "#"
        context["ex_type_{}".format(set_ex_type_list_transfer[type_index])] = set_ex_type_list[type_index]

# 升温曲线
img_index = 1
for ex_num in range(1, set_ex_num + 1):
    img_path = 'data/curve-{}-EN/{}.png'.format(save_img_time, index)
    context["curve_img_{}_heating_{}".format(ex_num, img_index)] = InlineImage(
        doc,
        image_descriptor=img_path,
        width=Cm(16.26),
        height=Cm(9.15))
    img_index += 1
    index += 1

# ---------------------------------------------------------------
# 实验运行记录表格
# set_ex_type_list_transfer = ['545_10', 'minus_545_10', '545_100', 'minus_545_100', '1000_10', 'minus_1000_10',
#                              '1000_100',
#                              'minus_1000_100',
#                              '1500_100', 'minus_1500_100', '1900_100', 'minus_1900_100']

for ex_num in range(set_ex_num):
    ex_num_dict = ex_num + 1
    context["ex_static_id_" + str(ex_num_dict)] = str(ex_num_dict) + "#"
    context["ex_time_" + str(ex_num_dict)] = set_ex_time[ex_num]
    context["ex_note_" + str(ex_num_dict)] = set_ex_note[ex_num]
    for index in range(len(set_ex_type_list_transfer)):
        context["ex_type_" + set_ex_type_list_transfer[index]] = set_ex_type_list[index]
        context["ex_phenomenon_" + str(ex_num_dict) + "_" + set_ex_type_list_transfer[index]] = set_ex_phenomenon[
            ex_num]


# ---------------------------------------------------------------
# 目录更新
def toc_main(docx_file):
    script_dir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
    file_name = docx_file
    file_path = os.path.join(script_dir, file_name)
    update_toc(file_path)


def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    try:
        doc = word.Documents.Open(docx_file)
        for contentCount in range(1, doc.TablesOfContents.Count + 1):
            doc.TablesOfContents(contentCount).Update()
        for FigureCount in range(1, doc.TablesOfFigures.Count + 1):
            doc.TablesOfFigures(FigureCount).Update()
        doc.Close(SaveChanges=True)
    finally:
        word.Quit()


toc_main(doc_temp_path)
# ---------------------------------------------------------------
doc.render(context)
doc.save(save_path)
